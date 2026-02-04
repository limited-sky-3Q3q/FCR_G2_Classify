# -*- coding: utf-8 -*-
"""
全部8种模型性能比较（雷达图）
包含：SGD、KNN、Neural Network (MLP)、SVM、Gradient Boosting、Random Forest、Extra Trees、Logistic Regression
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.model_selection import cross_val_predict, StratifiedKFold
from sklearn.preprocessing import StandardScaler, RobustScaler, PolynomialFeatures
from sklearn.neighbors import KNeighborsClassifier
from sklearn.linear_model import SGDClassifier, LogisticRegression
from sklearn.svm import SVC
from sklearn.ensemble import GradientBoostingClassifier, RandomForestClassifier, ExtraTreesClassifier
from sklearn.neural_network import MLPClassifier
from sklearn.metrics import confusion_matrix, roc_auc_score, roc_curve, brier_score_loss
import warnings
warnings.filterwarnings('ignore')

# Set English fonts
plt.rcParams['font.family'] = 'Arial'
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.size'] = 11

print("="*80)
print("全部8种模型性能比较（雷达图）")
print("="*80)

# 1. 数据加载
print("\n【1. 数据加载】")
data = pd.read_csv('CARCT1_FG2.csv', encoding='gbk')
data = data.dropna(subset=['FCR_G2', 'ID'])

exclude_cols = ['ID', 'FCR_G2', 'time', 'Group',
                'FCR1', 'FCR2', 'FCR3', 'FCR4',
                'PHQ9.1', 'PHQ9.2', 'PHQ9.3', 'PHQ9.4',
                'GAD7.1', 'GAD7.2', 'GAD7.3', 'GAD7.4',
                'ISI1', 'ISI2', 'ISI3', 'ISI4']

feature_cols = [col for col in data.columns if col not in exclude_cols]
X = data[feature_cols].copy()
y = data['FCR_G2'].copy().astype(int)

if 'FCR7_0' in X.columns:
    X = X.drop('FCR7_0', axis=1)

print(f"样本数: {len(data)}")
print(f"类别分布: 类别1={(y==1).sum()}, 类别2={(y==2).sum()}")

# 2. 准备不同模型的数据
print("\n【2. 准备模型数据】")

# SGD模型的8个特征
sgd_features = ['GAD7_0', 'Place', 'Education', 'Cohabitant',
                'Emotional_support', 'Recognition_disease_severity',
                'Knowingduration_of_cancer', 'Stress']
X_sgd = X[sgd_features]
X_sgd_scaled = RobustScaler().fit_transform(X_sgd)

# 其他模型的12个特征
knn_features = ['GAD7_0', 'PHQ9_0', 'ISI_0', 'TCSQ_NC',
                'Emotional_support', 'Stress', 'Exercise',
                'Education', 'Knowingduration_of_cancer',
                'Recognition_disease_severity', 'QLQBR23', 'TCSQ_PC']
X_knn = X[knn_features]
X_knn_scaled = StandardScaler().fit_transform(X_knn)

# Logistic Regression - RFE 12个特征 + 多项式特征
logreg_features = ['PHQ9_0', 'GAD7_0', 'ISI_0', 'TCSQ_PC', 'TCSQ_NC',
                   'CTQ', 'QLQBR23', 'Age', 'Education',
                   'Emotional_support', 'Recognition_disease_severity', 'Stress']
X_logreg = X[logreg_features]
poly = PolynomialFeatures(degree=2, include_bias=False)
X_logreg_poly = poly.fit_transform(X_logreg)
X_logreg_scaled = StandardScaler().fit_transform(X_logreg_poly)

# 3. 设置交叉验证
print("\n【3. 设置交叉验证】")

skf = StratifiedKFold(n_splits=8, shuffle=True, random_state=42)
print(f"使用 {skf.n_splits} 折交叉验证")

# 4. 计算各模型性能指标
print("\n【4. 计算性能指标】")

def calculate_metrics_with_ci(estimator_class, estimator_params, X_scaled, model_name, dataset_type="Validation"):
    """计算性能指标和95%置信区间 - 使用交叉验证在每一折重新训练模型"""
    print(f"  计算{model_name} {dataset_type}集性能...")

    # 计算每折的指标
    acc_list, prec_list, rec_list, f1_list, spec_list, auc_list = [], [], [], [], [], []
    ppv_list, npv_list, brier_list, sensitivity_list = [], [], [], []

    for train_idx, test_idx in skf.split(X_scaled, y):
        X_train, X_test = X_scaled[train_idx], X_scaled[test_idx]
        y_train, y_test = y.iloc[train_idx], y.iloc[test_idx]

        # 在每一折中重新创建并训练模型
        model = estimator_class(**estimator_params)
        model.fit(X_train, y_train)

        # 训练集预测
        if dataset_type == "Training":
            y_pred = model.predict(X_train)
            y_true = y_train
            # 训练集AUC需要预测概率
            # 训练集Brier Score需要预测概率
            try:
                if hasattr(model, 'predict_proba'):
                    y_proba = model.predict_proba(X_train)[:, 1]
                    # 处理标签为1和2的情况
                    y_true_auc = (y_true == 2).astype(int)
                    auc = roc_auc_score(y_true_auc, y_proba) if len(np.unique(y_true_auc)) > 1 else np.nan
                    brier = brier_score_loss(y_true_auc, y_proba)
                else:
                    auc = np.nan
                    brier = np.nan
            except:
                auc = np.nan
                brier = np.nan
        # 验证集预测
        else:
            y_pred = model.predict(X_test)
            y_true = y_test
            # 验证集AUC需要预测概率
            # 验证集Brier Score需要预测概率
            try:
                if hasattr(model, 'predict_proba'):
                    y_proba = model.predict_proba(X_test)[:, 1]
                    # 处理标签为1和2的情况
                    y_true_auc = (y_true == 2).astype(int)
                    auc = roc_auc_score(y_true_auc, y_proba) if len(np.unique(y_true_auc)) > 1 else np.nan
                    brier = brier_score_loss(y_true_auc, y_proba)
                else:
                    auc = np.nan
                    brier = np.nan
            except:
                auc = np.nan
                brier = np.nan

        cm = confusion_matrix(y_true, y_pred)
        if cm.size == 4:
            tn, fp, fn, tp = cm.ravel()
        else:
            tn, fp, fn, tp = 0, 0, 0, 0

        # 基本指标
        accuracy = (tp + tn) / (tp + tn + fp + fn) if (tp + tn + fp + fn) > 0 else 0
        precision = tp / (tp + fp) if (tp + fp) > 0 else 0
        recall = tp / (tp + fn) if (tp + fn) > 0 else 0
        specificity = tn / (tn + fp) if (tn + fp) > 0 else 0
        f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0

        # PPV (Positive Predictive Value) = Precision
        ppv = precision

        # NPV (Negative Predictive Value) = TN / (TN + FN)
        npv = tn / (tn + fn) if (tn + fn) > 0 else 0

        # Sensitivity = Recall
        sensitivity = recall

        acc_list.append(accuracy)
        prec_list.append(precision)
        rec_list.append(recall)
        f1_list.append(f1)
        spec_list.append(specificity)
        ppv_list.append(ppv)
        npv_list.append(npv)
        brier_list.append(brier)
        sensitivity_list.append(sensitivity)
        if not np.isnan(auc):
            auc_list.append(auc)

    # 计算平均值
    accuracy = np.mean(acc_list)
    precision = np.mean(prec_list)
    recall = np.mean(rec_list)
    f1 = np.mean(f1_list)
    specificity = np.mean(spec_list)
    ppv = np.mean(ppv_list)
    npv = np.mean(npv_list)
    brier = np.mean(brier_list)
    sensitivity = np.mean(sensitivity_list)
    auc = np.mean(auc_list) if len(auc_list) > 0 else np.nan

    # 计算标准差和95%置信区间
    acc_std = np.std(acc_list, ddof=1)
    prec_std = np.std(prec_list, ddof=1)
    rec_std = np.std(rec_list, ddof=1)
    f1_std = np.std(f1_list, ddof=1)
    spec_std = np.std(spec_list, ddof=1)
    auc_std = np.std(auc_list, ddof=1) if len(auc_list) > 1 else 0

    # 95% CI计算公式: mean ± 1.96 * std / sqrt(n)
    n_folds = len(acc_list)
    ci_multiplier = 1.96 / np.sqrt(n_folds)

    acc_ci = (accuracy - ci_multiplier * acc_std, accuracy + ci_multiplier * acc_std)
    auc_ci = (auc - ci_multiplier * auc_std, auc + ci_multiplier * auc_std) if not np.isnan(auc) else (np.nan, np.nan)

    print(f"    Accuracy={accuracy:.4f}±{acc_std:.4f}, PPV={ppv:.4f}±{prec_std:.4f}, Sensitivity={sensitivity:.4f}±{rec_std:.4f}, Brier={brier:.4f}, AUC={auc:.4f}")

    return {
        'accuracy': accuracy,
        'ppv': ppv,
        'npv': npv,
        'sensitivity': sensitivity,
        'specificity': specificity,
        'brier': brier,
        'f1': f1,
        'auc': auc,
        'accuracy_std': acc_std,
        'ppv_std': prec_std,
        'sensitivity_std': rec_std,
        'f1_std': f1_std,
        'auc_std': auc_std,
        'accuracy_ci': acc_ci,
        'auc_ci': auc_ci
    }

# 定义各模型的类和参数
sgd_params = {'loss': 'log_loss', 'penalty': 'elasticnet', 'alpha': 0.0002,
              'l1_ratio': 0.1, 'max_iter': 3000, 'class_weight': 'balanced', 'random_state': 42}
knn_params = {'n_neighbors': 6, 'weights': 'uniform', 'metric': 'manhattan', 'p': 1}
mlp_params = {'hidden_layer_sizes': (100, 50), 'activation': 'tanh', 'alpha': 0.1,
              'learning_rate_init': 0.1, 'max_iter': 500, 'batch_size': 'auto', 'random_state': 42}
svm_params = {'C': 0.1, 'kernel': 'sigmoid', 'gamma': 0.5, 'degree': 2,
              'class_weight': 'balanced', 'probability': True, 'random_state': 42}
gb_params = {'n_estimators': 50, 'learning_rate': 0.05, 'max_depth': 3,
            'min_samples_split': 5, 'min_samples_leaf': 1, 'subsample': 1.0, 'random_state': 42}
rf_params = {'n_estimators': 100, 'max_depth': None, 'max_features': None,
            'min_samples_leaf': 1, 'min_samples_split': 7, 'class_weight': 'balanced', 'random_state': 42}
et_params = {'n_estimators': 100, 'max_depth': None, 'max_features': 'sqrt',
            'min_samples_leaf': 4, 'min_samples_split': 2, 'class_weight': 'balanced', 'random_state': 42}
lr_params = {'C': 0.001, 'penalty': 'l2', 'solver': 'liblinear', 'max_iter': 1000, 'random_state': 42}

# 计算验证集性能
sgd_val = calculate_metrics_with_ci(SGDClassifier, sgd_params, X_sgd_scaled, "SGD", "Validation")
knn_val = calculate_metrics_with_ci(KNeighborsClassifier, knn_params, X_knn_scaled, "KNN", "Validation")
mlp_val = calculate_metrics_with_ci(MLPClassifier, mlp_params, X_knn_scaled, "MLP", "Validation")
svm_val = calculate_metrics_with_ci(SVC, svm_params, X_knn_scaled, "SVM", "Validation")
gb_val = calculate_metrics_with_ci(GradientBoostingClassifier, gb_params, X_knn_scaled, "Gradient Boosting", "Validation")
rf_val = calculate_metrics_with_ci(RandomForestClassifier, rf_params, X_knn_scaled, "Random Forest", "Validation")
et_val = calculate_metrics_with_ci(ExtraTreesClassifier, et_params, X_knn_scaled, "Extra Trees", "Validation")
lr_val = calculate_metrics_with_ci(LogisticRegression, lr_params, X_logreg_scaled, "Logistic Regression", "Validation")

# 计算训练集性能
sgd_train = calculate_metrics_with_ci(SGDClassifier, sgd_params, X_sgd_scaled, "SGD", "Training")
knn_train = calculate_metrics_with_ci(KNeighborsClassifier, knn_params, X_knn_scaled, "KNN", "Training")
mlp_train = calculate_metrics_with_ci(MLPClassifier, mlp_params, X_knn_scaled, "MLP", "Training")
svm_train = calculate_metrics_with_ci(SVC, svm_params, X_knn_scaled, "SVM", "Training")
gb_train = calculate_metrics_with_ci(GradientBoostingClassifier, gb_params, X_knn_scaled, "Gradient Boosting", "Training")
rf_train = calculate_metrics_with_ci(RandomForestClassifier, rf_params, X_knn_scaled, "Random Forest", "Training")
et_train = calculate_metrics_with_ci(ExtraTreesClassifier, et_params, X_knn_scaled, "Extra Trees", "Training")
lr_train = calculate_metrics_with_ci(LogisticRegression, lr_params, X_logreg_scaled, "Logistic Regression", "Training")

# 5. 生成雷达图
print("\n【5. 生成雷达图】")

categories = ['Accuracy', 'PPV', 'Sensitivity', 'F1-Score', 'Specificity']

# 确保所有指标值都是有效的，缺失的设置为0
def safe_value(val):
    """确保值为有效数字，缺失则返回0"""
    if val is None or np.isnan(val) or np.isinf(val):
        return 0.0
    return val

sgd_values = [safe_value(sgd_val['accuracy']), safe_value(sgd_val['ppv']), safe_value(sgd_val['sensitivity']), safe_value(sgd_val['f1']), safe_value(sgd_val['specificity'])]
knn_values = [safe_value(knn_val['accuracy']), safe_value(knn_val['ppv']), safe_value(knn_val['sensitivity']), safe_value(knn_val['f1']), safe_value(knn_val['specificity'])]
mlp_values = [safe_value(mlp_val['accuracy']), safe_value(mlp_val['ppv']), safe_value(mlp_val['sensitivity']), safe_value(mlp_val['f1']), safe_value(mlp_val['specificity'])]
svm_values = [safe_value(svm_val['accuracy']), safe_value(svm_val['ppv']), safe_value(svm_val['sensitivity']), safe_value(svm_val['f1']), safe_value(svm_val['specificity'])]
gb_values = [safe_value(gb_val['accuracy']), safe_value(gb_val['ppv']), safe_value(gb_val['sensitivity']), safe_value(gb_val['f1']), safe_value(gb_val['specificity'])]
rf_values = [safe_value(rf_val['accuracy']), safe_value(rf_val['ppv']), safe_value(rf_val['sensitivity']), safe_value(rf_val['f1']), safe_value(rf_val['specificity'])]
et_values = [safe_value(et_val['accuracy']), safe_value(et_val['ppv']), safe_value(et_val['sensitivity']), safe_value(et_val['f1']), safe_value(et_val['specificity'])]
lr_values = [safe_value(lr_val['accuracy']), safe_value(lr_val['ppv']), safe_value(lr_val['sensitivity']), safe_value(lr_val['f1']), safe_value(lr_val['specificity'])]

# Unified color scheme for all 8 models (consistent across all charts)
COLORS = {
    'SGD': '#E74C3C',           # Red
    'KNN': '#3498DB',           # Blue
    'MLP': '#2ECC71',           # Green
    'SVM': '#F39C12',           # Orange
    'GB': '#9B59B6',            # Purple
    'RF': '#1ABC9C',            # Teal
    'ET': '#E67E22',            # Dark Orange
    'LR': '#16A085',            # Dark Teal
}

colors = [COLORS['SGD'], COLORS['KNN'], COLORS['MLP'], COLORS['SVM'],
          COLORS['GB'], COLORS['RF'], COLORS['ET'], COLORS['LR']]

# 使用平方变换来进一步突出高分数值的差异
def transform_values(values):
    """使用平方变换，使高值区域拉伸更大"""
    # 将值从[0, 1]范围变换为平方，极大拉伸高值区域
    return values ** 2

# 定义原始值数组
sgd_values_raw = np.array([safe_value(sgd_val['accuracy']), safe_value(sgd_val['ppv']), safe_value(sgd_val['sensitivity']), safe_value(sgd_val['f1']), safe_value(sgd_val['specificity'])])
knn_values_raw = np.array([safe_value(knn_val['accuracy']), safe_value(knn_val['ppv']), safe_value(knn_val['sensitivity']), safe_value(knn_val['f1']), safe_value(knn_val['specificity'])])
mlp_values_raw = np.array([safe_value(mlp_val['accuracy']), safe_value(mlp_val['ppv']), safe_value(mlp_val['sensitivity']), safe_value(mlp_val['f1']), safe_value(mlp_val['specificity'])])
svm_values_raw = np.array([safe_value(svm_val['accuracy']), safe_value(svm_val['ppv']), safe_value(svm_val['sensitivity']), safe_value(svm_val['f1']), safe_value(svm_val['specificity'])])
gb_values_raw = np.array([safe_value(gb_val['accuracy']), safe_value(gb_val['ppv']), safe_value(gb_val['sensitivity']), safe_value(gb_val['f1']), safe_value(gb_val['specificity'])])
rf_values_raw = np.array([safe_value(rf_val['accuracy']), safe_value(rf_val['ppv']), safe_value(rf_val['sensitivity']), safe_value(rf_val['f1']), safe_value(rf_val['specificity'])])
et_values_raw = np.array([safe_value(et_val['accuracy']), safe_value(et_val['ppv']), safe_value(et_val['sensitivity']), safe_value(et_val['f1']), safe_value(et_val['specificity'])])
lr_values_raw = np.array([safe_value(lr_val['accuracy']), safe_value(lr_val['ppv']), safe_value(lr_val['sensitivity']), safe_value(lr_val['f1']), safe_value(lr_val['specificity'])])

# Transform all values to non-linear scale
sgd_values_t = list(transform_values(sgd_values_raw)) + [transform_values(sgd_values_raw)[0]]
knn_values_t = list(transform_values(knn_values_raw)) + [transform_values(knn_values_raw)[0]]
mlp_values_t = list(transform_values(mlp_values_raw)) + [transform_values(mlp_values_raw)[0]]
svm_values_t = list(transform_values(svm_values_raw)) + [transform_values(svm_values_raw)[0]]
gb_values_t = list(transform_values(gb_values_raw)) + [transform_values(gb_values_raw)[0]]
rf_values_t = list(transform_values(rf_values_raw)) + [transform_values(rf_values_raw)[0]]
et_values_t = list(transform_values(et_values_raw)) + [transform_values(et_values_raw)[0]]
lr_values_t = list(transform_values(lr_values_raw)) + [transform_values(lr_values_raw)[0]]

# 设置雷达图
fig, ax = plt.subplots(figsize=(14, 11), subplot_kw=dict(projection='polar'))

angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
angles += angles[:1]

# Plot all 8 models with unified colors
markers = ['o', 's', '^', 'D', 'v', 'p', '*', 'h']
linestyles = ['-', '-', '-', '-', '-', '-', '-', '-']

labels = [f'SGD\n({sgd_val["accuracy"]*100:.2f}%)',
          f'KNN\n({knn_val["accuracy"]*100:.2f}%)',
          f'Neural Net\n({mlp_val["accuracy"]*100:.2f}%)',
          f'SVM\n({svm_val["accuracy"]*100:.2f}%)',
          f'Gradient Boost\n({gb_val["accuracy"]*100:.2f}%)',
          f'Random Forest\n({rf_val["accuracy"]*100:.2f}%)',
          f'Extra Trees\n({et_val["accuracy"]*100:.2f}%)',
          f'Logistic Reg\n({lr_val["accuracy"]*100:.2f}%)']

# 先绘制所有模型的填充
for i, (values, color) in enumerate(zip(
    [sgd_values_t, knn_values_t, mlp_values_t, svm_values_t, gb_values_t, rf_values_t, et_values_t, lr_values_t],
    colors
)):
    ax.fill(angles, values, alpha=0.08, color=color, zorder=1)

# 再绘制所有模型的轮廓（不填充）- 放在上层
for i, (values, color, marker, label) in enumerate(zip(
    [sgd_values_t, knn_values_t, mlp_values_t, svm_values_t, gb_values_t, rf_values_t, et_values_t, lr_values_t],
    colors, markers, labels
)):
    ax.plot(angles, values, marker=marker, linewidth=3.5, label=label,
            color=color, markersize=10, linestyle=linestyles[i], zorder=5)

# 设置雷达图属性 - 使用更窄的范围以突出差异
ax.set_xticks(angles[:-1])
ax.set_xticklabels(categories, fontsize=14, fontweight='bold', zorder=100)

# 设置Y轴范围 - 基于变换后的值（平方变换）
ax.set_ylim(0.36, 1.0)  # 变换后的范围：0.60^2 = 0.36 到 1.00^2 = 1.0

# 设置Y轴刻度 - 显示原始值而非变换值
ytick_original = [0.60, 0.70, 0.80, 0.90, 1.00]
ytick_transformed = transform_values(np.array(ytick_original))

ax.set_yticks(ytick_transformed)
ax.set_yticklabels([f'{v:.0%}' for v in ytick_original], fontsize=11, fontweight='bold')

# 将Y轴标签设置到最上层
for tick in ax.get_yticklabels():
    tick.set_zorder(1000)

# 将网格线设置在较低层级
ax.grid(True, alpha=0.4, linestyle='--', linewidth=1.5, zorder=2)

# Add title
ax.set_title('Performance Comparison of 8 Models (Radar Chart)\nFCR Trajectory Prediction in 52 Breast Cancer Patients',
             fontsize=16, fontweight='bold', pad=35, y=1.08)

# 设置图例 - 放在右侧
ax.legend(loc='center left', bbox_to_anchor=(1.20, 0.5), fontsize=11,
          frameon=True, fancybox=True, shadow=True)

plt.tight_layout()
plt.savefig('全部8模型性能比较（雷达图）.png', dpi=300, bbox_inches='tight')
plt.savefig('全部8模型性能比较（雷达图）.pdf', dpi=300, bbox_inches='tight')
plt.close()

print("\n已保存: 全部8模型性能比较（雷达图）.png")
print("\n已保存: 全部8模型性能比较（雷达图）.pdf")

# 6. 生成训练集和验证集的完整三线表
print("\n【6. 生成训练集和验证集三线表】")

def create_table_data(metrics_dict, dataset_name="Validation"):
    """创建表格数据 - 只对AUC和Accuracy显示95%CI"""
    table_data = [
        ['SGDClassifier',
         f'{metrics_dict["SGD"]["accuracy"]:.3f} [{metrics_dict["SGD"]["accuracy_ci"][0]:.3f}, {metrics_dict["SGD"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["SGD"]["auc"]:.3f} [{metrics_dict["SGD"]["auc_ci"][0]:.3f}, {metrics_dict["SGD"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["SGD"]["auc"]) else 'N/A',
         f'{metrics_dict["SGD"]["ppv"]:.3f}',
         f'{metrics_dict["SGD"]["npv"]:.3f}',
         f'{metrics_dict["SGD"]["sensitivity"]:.3f}',
         f'{metrics_dict["SGD"]["specificity"]:.3f}',
         f'{metrics_dict["SGD"]["brier"]:.3f}',
         f'{metrics_dict["SGD"]["f1"]:.3f}'],
        ['KNN',
         f'{metrics_dict["KNN"]["accuracy"]:.3f} [{metrics_dict["KNN"]["accuracy_ci"][0]:.3f}, {metrics_dict["KNN"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["KNN"]["auc"]:.3f} [{metrics_dict["KNN"]["auc_ci"][0]:.3f}, {metrics_dict["KNN"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["KNN"]["auc"]) else 'N/A',
         f'{metrics_dict["KNN"]["ppv"]:.3f}',
         f'{metrics_dict["KNN"]["npv"]:.3f}',
         f'{metrics_dict["KNN"]["sensitivity"]:.3f}',
         f'{metrics_dict["KNN"]["specificity"]:.3f}',
         f'{metrics_dict["KNN"]["brier"]:.3f}',
         f'{metrics_dict["KNN"]["f1"]:.3f}'],
        ['Neural Network (MLP)',
         f'{metrics_dict["MLP"]["accuracy"]:.3f} [{metrics_dict["MLP"]["accuracy_ci"][0]:.3f}, {metrics_dict["MLP"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["MLP"]["auc"]:.3f} [{metrics_dict["MLP"]["auc_ci"][0]:.3f}, {metrics_dict["MLP"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["MLP"]["auc"]) else 'N/A',
         f'{metrics_dict["MLP"]["ppv"]:.3f}',
         f'{metrics_dict["MLP"]["npv"]:.3f}',
         f'{metrics_dict["MLP"]["sensitivity"]:.3f}',
         f'{metrics_dict["MLP"]["specificity"]:.3f}',
         f'{metrics_dict["MLP"]["brier"]:.3f}',
         f'{metrics_dict["MLP"]["f1"]:.3f}'],
        ['SVM',
         f'{metrics_dict["SVM"]["accuracy"]:.3f} [{metrics_dict["SVM"]["accuracy_ci"][0]:.3f}, {metrics_dict["SVM"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["SVM"]["auc"]:.3f} [{metrics_dict["SVM"]["auc_ci"][0]:.3f}, {metrics_dict["SVM"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["SVM"]["auc"]) else 'N/A',
         f'{metrics_dict["SVM"]["ppv"]:.3f}',
         f'{metrics_dict["SVM"]["npv"]:.3f}',
         f'{metrics_dict["SVM"]["sensitivity"]:.3f}',
         f'{metrics_dict["SVM"]["specificity"]:.3f}',
         f'{metrics_dict["SVM"]["brier"]:.3f}',
         f'{metrics_dict["SVM"]["f1"]:.3f}'],
        ['Gradient Boosting',
         f'{metrics_dict["GB"]["accuracy"]:.3f} [{metrics_dict["GB"]["accuracy_ci"][0]:.3f}, {metrics_dict["GB"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["GB"]["auc"]:.3f} [{metrics_dict["GB"]["auc_ci"][0]:.3f}, {metrics_dict["GB"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["GB"]["auc"]) else 'N/A',
         f'{metrics_dict["GB"]["ppv"]:.3f}',
         f'{metrics_dict["GB"]["npv"]:.3f}',
         f'{metrics_dict["GB"]["sensitivity"]:.3f}',
         f'{metrics_dict["GB"]["specificity"]:.3f}',
         f'{metrics_dict["GB"]["brier"]:.3f}',
         f'{metrics_dict["GB"]["f1"]:.3f}'],
        ['Random Forest',
         f'{metrics_dict["RF"]["accuracy"]:.3f} [{metrics_dict["RF"]["accuracy_ci"][0]:.3f}, {metrics_dict["RF"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["RF"]["auc"]:.3f} [{metrics_dict["RF"]["auc_ci"][0]:.3f}, {metrics_dict["RF"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["RF"]["auc"]) else 'N/A',
         f'{metrics_dict["RF"]["ppv"]:.3f}',
         f'{metrics_dict["RF"]["npv"]:.3f}',
         f'{metrics_dict["RF"]["sensitivity"]:.3f}',
         f'{metrics_dict["RF"]["specificity"]:.3f}',
         f'{metrics_dict["RF"]["brier"]:.3f}',
         f'{metrics_dict["RF"]["f1"]:.3f}'],
        ['Extra Trees',
         f'{metrics_dict["ET"]["accuracy"]:.3f} [{metrics_dict["ET"]["accuracy_ci"][0]:.3f}, {metrics_dict["ET"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["ET"]["auc"]:.3f} [{metrics_dict["ET"]["auc_ci"][0]:.3f}, {metrics_dict["ET"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["ET"]["auc"]) else 'N/A',
         f'{metrics_dict["ET"]["ppv"]:.3f}',
         f'{metrics_dict["ET"]["npv"]:.3f}',
         f'{metrics_dict["ET"]["sensitivity"]:.3f}',
         f'{metrics_dict["ET"]["specificity"]:.3f}',
         f'{metrics_dict["ET"]["brier"]:.3f}',
         f'{metrics_dict["ET"]["f1"]:.3f}'],
        ['Logistic Regression',
         f'{metrics_dict["LR"]["accuracy"]:.3f} [{metrics_dict["LR"]["accuracy_ci"][0]:.3f}, {metrics_dict["LR"]["accuracy_ci"][1]:.3f}]',
         f'{metrics_dict["LR"]["auc"]:.3f} [{metrics_dict["LR"]["auc_ci"][0]:.3f}, {metrics_dict["LR"]["auc_ci"][1]:.3f}]' if not np.isnan(metrics_dict["LR"]["auc"]) else 'N/A',
         f'{metrics_dict["LR"]["ppv"]:.3f}',
         f'{metrics_dict["LR"]["npv"]:.3f}',
         f'{metrics_dict["LR"]["sensitivity"]:.3f}',
         f'{metrics_dict["LR"]["specificity"]:.3f}',
         f'{metrics_dict["LR"]["brier"]:.3f}',
         f'{metrics_dict["LR"]["f1"]:.3f}'],
    ]

    df_table = pd.DataFrame(table_data,
                        columns=['Model',
                                'Accuracy (95% CI)',
                                'AUC (95% CI)',
                                'PPV',
                                'NPV',
                                'Sensitivity',
                                'Specificity',
                                'Brier',
                                'F1'])
    return df_table

def save_table_and_plot(df_table, dataset_name="Validation"):
    """保存表格数据和生成图片"""
    # 保存CSV
    csv_filename = f'全部8模型性能对比三线表_{dataset_name}.csv'
    df_table.to_csv(csv_filename, index=False, encoding='utf-8-sig')
    print(f"已保存: {csv_filename}")

    # 保存Word文档
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 添加标题
        title = doc.add_heading(f'Performance Comparison of 8 Models on {dataset_name} Set (Three-Line Table)', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph('FCR Trajectory Prediction in 52 Breast Cancer Patients')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加表格
        table = doc.add_table(rows=len(df_table) + 1, cols=len(df_table.columns))
        table.style = 'Table Grid'

        # 设置表头
        for i, col_name in enumerate(df_table.columns):
            cell = table.rows[0].cells[i]
            cell.text = col_name
            # 设置表头格式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            cell.background_color = '#2E86AB'
            # 设置表头背景色
            from docx.oxml.shared import OxmlElement, qn
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '2E86AB')
            cell._element.get_or_add_tcPr().append(shading_elm)

        # 设置表格内容
        for i, row_data in enumerate(df_table.values):
            for j, cell_value in enumerate(row_data):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(cell_value)

                # 设置单元格格式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)

                # 设置斑马纹背景
                row_colors = ['#D6EAF8', '#EBF5FB']
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), row_colors[i % 2][1:])
                cell._element.get_or_add_tcPr().append(shading_elm)

                # 第一列（模型名称）使用彩色背景
                if j == 0:
                    model_name = row_data
                    color_map = {
                        'SGD': '#E74C3C',
                        'KNN': '#3498DB',
                        'Neural': '#2ECC71',
                        'MLP': '#2ECC71',
                        'SVM': '#F39C12',
                        'Gradient': '#9B59B6',
                        'Random': '#1ABC9C',
                        'Extra': '#E67E22',
                        'Logistic': '#16A085'
                    }
                    color_code = '#FFFFFF'
                    for key, color in color_map.items():
                        if key in model_name:
                            color_code = color
                            break

                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), color_code[1:])
                    cell._element.get_or_add_tcPr().append(shading_elm)

                    # 设置模型名称为白色粗体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)

        # 保存Word文件
        docx_filename = f'全部8模型性能对比三线表_{dataset_name}.docx'
        doc.save(docx_filename)
        print(f"已保存: {docx_filename}")

    except ImportError:
        print("注意: 未安装python-docx库，跳过Word文档保存")
        print("      可使用命令安装: pip install python-docx")
    except Exception as e:
        print(f"保存Word文档时出错: {str(e)}")

    # 生成表格图片
    fig, ax = plt.subplots(figsize=(24, 10))
    ax.axis('tight')
    ax.axis('off')

    table = ax.table(cellText=df_table.values,
                  colLabels=df_table.columns,
                  cellLoc='center',
                  loc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 3.5)

    # 设置表头样式
    for i in range(len(df_table.columns)):
        table[(0, i)].set_facecolor('#2E86AB')
        table[(0, i)].set_text_props(weight='bold', color='white', fontsize=10)

    # 设置斑马纹
    colors_rows = ['#D6EAF8', '#EBF5FB']
    for i in range(1, len(df_table) + 1):
        for j in range(len(df_table.columns)):
            table[(i, j)].set_facecolor(colors_rows[(i-1) % 2])
        # Highlight model names with unified colors
        model_name = df_table.iloc[i-1, 0]
        if 'SGD' in model_name:
            table[(i, 0)].set_facecolor(COLORS['SGD'])
            table[(i, 0)].set_text_props(weight='bold', color='white')
        elif 'KNN' in model_name:
            table[(i, 0)].set_facecolor(COLORS['KNN'])
            table[(i, 0)].set_text_props(weight='bold', color='white')
        elif 'Neural' in model_name or 'MLP' in model_name:
            table[(i, 0)].set_facecolor(COLORS['MLP'])
            table[(i, 0)].set_text_props(weight='bold', color='white')
        elif 'SVM' in model_name:
            table[(i, 0)].set_facecolor(COLORS['SVM'])
            table[(i, 0)].set_text_props(weight='bold', color='white')
        elif 'Gradient' in model_name:
            table[(i, 0)].set_facecolor(COLORS['GB'])
            table[(i, 0)].set_text_props(weight='bold', color='white')
        elif 'Random' in model_name:
            table[(i, 0)].set_facecolor(COLORS['RF'])
            table[(i, 0)].set_text_props(weight='bold', color='white')
        elif 'Extra' in model_name:
            table[(i, 0)].set_facecolor(COLORS['ET'])
            table[(i, 0)].set_text_props(weight='bold', color='white')
        elif 'Logistic' in model_name:
            table[(i, 0)].set_facecolor(COLORS['LR'])
            table[(i, 0)].set_text_props(weight='bold', color='white')

    ax.set_title(f'Performance Comparison of 8 Models on {dataset_name} Set (Three-Line Table)\nFCR Trajectory Prediction in 52 Breast Cancer Patients',
                 fontsize=16, fontweight='bold', pad=20)

    plt.tight_layout()
    png_filename = f'全部8模型性能对比三线表_{dataset_name}.png'
    pdf_filename = f'全部8模型性能对比三线表_{dataset_name}.pdf'
    plt.savefig(png_filename, dpi=300, bbox_inches='tight')
    plt.savefig(pdf_filename, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"已保存: {png_filename}")
    print(f"已保存: {pdf_filename}")

# 准备验证集数据
validation_metrics = {
    'SGD': sgd_val,
    'KNN': knn_val,
    'MLP': mlp_val,
    'SVM': svm_val,
    'GB': gb_val,
    'RF': rf_val,
    'ET': et_val,
    'LR': lr_val
}

# 准备训练集数据
training_metrics = {
    'SGD': sgd_train,
    'KNN': knn_train,
    'MLP': mlp_train,
    'SVM': svm_train,
    'GB': gb_train,
    'RF': rf_train,
    'ET': et_train,
    'LR': lr_train
}

# 生成验证集三线表
print("\n生成验证集三线表...")
df_validation = create_table_data(validation_metrics, "Validation")
save_table_and_plot(df_validation, "Validation")

# 生成训练集三线表
print("\n生成训练集三线表...")
df_training = create_table_data(training_metrics, "Training")
save_table_and_plot(df_training, "Training")


# 8. Summary
print("\n" + "="*80)
print("全部8种模型性能比较完成")
print("="*80)

generated_files = [
    "全部8模型性能比较（雷达图）.png",
    "全部8模型性能比较（雷达图）.pdf",
    "全部8模型性能对比三线表_Validation.csv",
    "全部8模型性能对比三线表_Validation.png",
    "全部8模型性能对比三线表_Validation.pdf",
    "全部8模型性能对比三线表_Training.csv",
    "全部8模型性能对比三线表_Training.png",
    "全部8模型性能对比三线表_Training.pdf",
]

print("\n已生成的文件:")
for i, file in enumerate(generated_files, 1):
    print(f"  {i}. {file}")

print(f"\n总计: {len(generated_files)} 个文件")

# 性能排序总结 - 验证集
print("\n【验证集性能排序】")
print("\n按准确率排序:")
print(f"{'排名':<6} {'模型':<25} {'准确率':>12} {'标准差':>12}")
print("-" * 60)

rank_data_val = []
model_metrics = [
    ('SGD', sgd_val), ('KNN', knn_val), ('MLP', mlp_val), ('SVM', svm_val),
    ('GB', gb_val), ('RF', rf_val), ('ET', et_val), ('LR', lr_val)
]

for model_name, metrics in model_metrics:
    rank_data_val.append({
        'model': model_name,
        'accuracy': metrics['accuracy'],
        'std': metrics['accuracy_std']
    })

rank_data_val.sort(key=lambda x: x['accuracy'], reverse=True)

for i, item in enumerate(rank_data_val, 1):
    print(f"{i:<6} {item['model']:<25} {item['accuracy']*100:>8.2f}%     ±{item['std']*100:>6.2f}%")

# 性能差异分析 - 验证集
print("\n【验证集性能差异分析】")
print(f"\n最佳模型: {rank_data_val[0]['model']} (准确率: {rank_data_val[0]['accuracy']*100:.2f}%)")
print(f"最差模型: {rank_data_val[-1]['model']} (准确率: {rank_data_val[-1]['accuracy']*100:.2f}%)")
print(f"性能差距: {rank_data_val[0]['accuracy']*100 - rank_data_val[-1]['accuracy']*100:.2f}%")

# 性能排序总结 - 训练集
print("\n\n【训练集性能排序】")
print("\n按准确率排序:")
print(f"{'排名':<6} {'模型':<25} {'准确率':>12} {'标准差':>12}")
print("-" * 60)

rank_data_train = []
model_metrics_train = [
    ('SGD', sgd_train), ('KNN', knn_train), ('MLP', mlp_train), ('SVM', svm_train),
    ('GB', gb_train), ('RF', rf_train), ('ET', et_train), ('LR', lr_train)
]

for model_name, metrics in model_metrics_train:
    rank_data_train.append({
        'model': model_name,
        'accuracy': metrics['accuracy'],
        'std': metrics['accuracy_std']
    })

rank_data_train.sort(key=lambda x: x['accuracy'], reverse=True)

for i, item in enumerate(rank_data_train, 1):
    print(f"{i:<6} {item['model']:<25} {item['accuracy']*100:>8.2f}%     ±{item['std']*100:>6.2f}%")

# 性能差异分析 - 训练集
print("\n【训练集性能差异分析】")
print(f"\n最佳模型: {rank_data_train[0]['model']} (准确率: {rank_data_train[0]['accuracy']*100:.2f}%)")
print(f"最差模型: {rank_data_train[-1]['model']} (准确率: {rank_data_train[-1]['accuracy']*100:.2f}%)")
print(f"性能差距: {rank_data_train[0]['accuracy']*100 - rank_data_train[-1]['accuracy']*100:.2f}%")

print("\n" + "="*80)
print("完成")
print("="*80)
